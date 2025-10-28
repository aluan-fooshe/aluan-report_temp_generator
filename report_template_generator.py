from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


class Another_report:
    def __init__(self, name, professor, school_class, report_title, margin_inch):
        dt = datetime.datetime.today()
        final_dt = datetime.datetime(dt.year, dt.month, dt.day)

        self.name = name
        self.professor = professor
        self.school_class = school_class
        self.date = final_dt.strftime("%d %B %Y")
        self.report_title = report_title

        self.margin_inch = margin_inch
        self.doc = Document()
        self.output_report = ''.join(['_' if char == ' ' else char for char in self.report_title])

        if name == None:
            self.name = "Audrey Luan"
            self.last_name = self.name.split(" ")[1]

        if report_title == None:
            self.report_title = f"{self.name}'s Report"

        if margin_inch == None:
            self.margin_inch = 1

        print(f"{self.name=}\n{self.professor}\n{self.school_class=}\n"
              f"{self.date=}\n{self.report_title=}\n{self.output_report=}")

    def page_numbers(self):
        header = self.doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"\t\t{self.last_name} 1"

    def set_margins(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def heading(self):
        heading = [self.name, self.professor, self.school_class, self.date]
        for item in heading:
            self.doc.add_paragraph(item, style="Normal")

    def title_report(self, align=WD_ALIGN_PARAGRAPH.CENTER, fontsize=None, fontname="Times New Roman"):
        if fontsize is None:
            fontsize = 12
        title = self.doc.add_paragraph(self.report_title)
        title.alignment = align
        title_format = title.runs[0]
        title_format.font.size = Pt(fontsize)
        title_format.font.name = fontname

    def add_paragraph(self, paragraph):
        self.doc.add_paragraph(paragraph)
        self.doc.save(f"{self.output_report}.docx")

    def add_subsection(self, subsection):
        subsection1 = self.doc.add_paragraph(subsection)
        subsection1.paragraph_format.line_spacing = Inches(0.5)
        self.doc.save(f"{self.output_report}.docx")

    def correct_document_pckg(self):
        try:
            document = Document(foo='bar', baz=12)
            print(document['baz'])

        except:
            print("Old Document class is gone — this failed as expected. The docx library is expected.\n")

if __name__ == '__main__':

    # Now you can import Class1
    from Filelist_workbook import Excel_Filelist
    import sys

    print(sys.executable)
    print(sys.path)

    # Test it
    obj = Excel_Filelist()
    subsections_dict = obj.import_dictionary(filename="report_template_subsections.txt")

    ass = Another_report(
        name=None,
        professor="Stephen C. Petersen",
        school_class="ECE 129A - Capstone pt 1",
        #school_class="ECE 118 - Why this class sucks",
        #report_title="Mechatronics Burn Book",
        report_title="C.A.R.T. - Carry Assist Robotic Transport",
        margin_inch=None
    )
    ass.correct_document_pckg()

    """ Set 1-inch margins """
    ass.set_margins()
    ass.heading()
    ass.title_report(
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fontsize=None
    )

    """ Set page numbers on every lab report page; '<Last_name> 1' """
    ass.page_numbers()

    for key, question in subsections_dict.items():
        print(question)
        ass.add_paragraph(f"{question}")
        ass.add_subsection( str(input("answer:\n")) )

    ass.doc.save(f"{ass.output_report}.docx")